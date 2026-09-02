"""
Block 13 — Data logging and report generation.

Changes from the documented version:
  * every row carries an explicit PASS/FAIL verdict against the framework
    limits, rather than a table of numbers a reader has to compare by eye;
  * the report records bench conditions (instrument identities, timestamp,
    whether the run was simulated) — Phase 8 of the framework document asks
    for "equipment serial numbers/cal dates, ambient conditions" and a table
    of three numbers does not satisfy that;
  * a simulated run is stamped as such in the heading. A generated .docx
    that looks identical whether or not real hardware was attached is a
    liability, not a feature.
"""

from __future__ import annotations

import datetime as _dt
from pathlib import Path

import pandas as pd

from .analysis import load_regulation_percent, setpoint_error_percent


def summarise(df: pd.DataFrame, limits) -> list[dict]:
    """Build the parameter / spec / measured / verdict rows."""
    rows: list[dict] = []

    def add(param, spec, measured, passed):
        rows.append({"parameter": param, "spec": spec,
                     "measured": measured,
                     "verdict": "PASS" if passed else "FAIL"})

    if "vout_dmm" in df and len(df) >= 2:
        reg = load_regulation_percent(df["vout_dmm"], limits.v_nominal)
        add("Load regulation (0 to I_max)",
            f"< {limits.regulation_percent}%",
            f"{reg:.3f}%", reg <= limits.regulation_percent)

        worst = df["vout_dmm"].sub(limits.v_nominal).abs().max()
        err = setpoint_error_percent(limits.v_nominal + worst, limits.v_nominal)
        add("Worst setpoint error",
            f"< {limits.regulation_percent}%",
            f"{err:.3f}%", err <= limits.regulation_percent)

        add("Min Vout at full load",
            f"> {limits.v_nominal * (1 - limits.regulation_percent / 100):.3f} V",
            f"{df['vout_dmm'].min():.4f} V",
            df["vout_dmm"].min()
            >= limits.v_nominal * (1 - limits.regulation_percent / 100))

    if "efficiency" in df and df["efficiency"].notna().any():
        eff = df["efficiency"].dropna()
        add("Peak efficiency", "reference only", f"{eff.max():.2f}%", True)
        full = df.dropna(subset=["efficiency"])
        if not full.empty:
            at_max = full.loc[full["load_current_set"].idxmax()]
            add(f"Efficiency at {at_max['load_current_set']:.1f} A",
                f"> {limits.efficiency_percent}%",
                f"{at_max['efficiency']:.2f}%",
                at_max["efficiency"] >= limits.efficiency_percent)

    return rows


def generate_report(df: pd.DataFrame, limits, plot_paths: list[str] | None = None,
                    output_doc: str = "Test_Report.docx", *,
                    instrument_ids: dict | None = None,
                    simulated: bool = False,
                    notes: str | None = None) -> str:
    """Write a Word test report. Returns the output path."""
    from docx import Document          # noqa: PLC0415
    from docx.shared import Inches     # noqa: PLC0415

    doc = Document()
    title = "Power Supply Test Report"
    if simulated:
        title += " (SIMULATED RUN - NOT HARDWARE DATA)"
    doc.add_heading(title, 0)

    if simulated:
        p = doc.add_paragraph()
        p.add_run(
            "This report was generated against the software bench simulator. "
            "The numbers below exercise the analysis path; they are not "
            "measurements of physical hardware."
        ).bold = True

    doc.add_heading("Test conditions", level=1)
    cond = doc.add_table(rows=0, cols=2)
    cond.style = "Table Grid"
    conditions = {
        "Date (UTC)": _dt.datetime.now(_dt.timezone.utc).strftime("%Y-%m-%d %H:%M"),
        "Data points": str(len(df)),
        "Nominal output": f"{limits.v_nominal} V",
        "Mode": "Simulated" if simulated else "Hardware",
    }
    for k, v in {**conditions, **(instrument_ids or {})}.items():
        cells = cond.add_row().cells
        cells[0].text, cells[1].text = str(k), str(v)

    doc.add_heading("Results against specification", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    for i, h in enumerate(["Parameter", "Spec", "Measured", "Verdict"]):
        table.rows[0].cells[i].text = h

    rows = summarise(df, limits)
    for r in rows:
        cells = table.add_row().cells
        cells[0].text = r["parameter"]
        cells[1].text = r["spec"]
        cells[2].text = r["measured"]
        cells[3].text = r["verdict"]

    failures = [r for r in rows if r["verdict"] == "FAIL"]
    doc.add_paragraph()
    verdict = doc.add_paragraph()
    verdict.add_run(
        f"Overall: {'FAIL' if failures else 'PASS'} "
        f"({len(failures)} of {len(rows)} checks failed)"
    ).bold = True

    if notes:
        doc.add_heading("Notes and known limitations", level=1)
        doc.add_paragraph(notes)

    for path in plot_paths or []:
        if Path(path).exists():
            doc.add_page_break()
            doc.add_picture(path, width=Inches(6.0))

    doc.save(output_doc)
    return output_doc
